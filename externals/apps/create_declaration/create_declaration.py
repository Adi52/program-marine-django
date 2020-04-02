from docx import Document


class CreateDeclaration:
    def __init__(self, parking_place, date, yacht, owner_details, commissioning_body, parking_peroid, fee, chip_card,
                 media):
        self.document = Document('externals/apps/create_declaration/deklaracja.docx')
        self.parking_peroid = parking_peroid
        self.parking_place = parking_place
        self.date = date
        self.yacht = yacht
        self.owner_details = owner_details
        self.commissioning_body = commissioning_body
        self.fee = fee
        self.chip_card = chip_card
        self.media = media

    def create_document(self):
        for paragraph in self.document.paragraphs:
            if 'Miejsce postojowe' in paragraph.text:
                paragraph.text = 'Miejsce postojowe {}                                                               ' \
                                 '                                        ' \
                                 'Data {}'.format(self.parking_place, self.date)

        self.document.add_paragraph('\n')
        self.document.add_paragraph('Nazwa jachtu: ', style='List Number').add_run(self.yacht['name']).bold = True
        self.document.add_paragraph('Nr rejestracyjny: ', style='List Number').add_run(
            self.yacht['registration_number']).bold = True
        self.document.add_paragraph('Port macierzysty: ', style='List Number').add_run(
            self.yacht['home_port']).bold = True
        self.document.add_paragraph('Dane jachtu: ', style='List Number')
        p = self.document.add_paragraph('            Długość: ')
        p.add_run(str(self.yacht['length'])).bold = True
        p.add_run(' m')
        p.add_run('             Szerokość: ')
        p.add_run(str(self.yacht['width'])).bold = True
        p.add_run(' m')
        self.document.add_paragraph('Typ jachtu: ', style='List Number').add_run(self.yacht['ytype']).bold = True
        self.document.add_paragraph('Dane właściciela* lub użytkownika jachtu*: ', style='List Number')
        self.document.add_paragraph('            - imię i nazwisko: ').add_run(self.owner_details['name']).bold = True
        self.document.add_paragraph('            - adres : ').add_run(self.owner_details['address']).bold = True
        self.document.add_paragraph('Dane armatora: ', style='List Number')
        self.document.add_paragraph('            - imię i nazwisko: ').add_run(self.owner_details['name']).bold = True
        self.document.add_paragraph('            - adres : ').add_run(self.owner_details['address']).bold = True
        self.document.add_paragraph('Podmiot zlecający, podpisujący umowę na postój jachtu: ', style='List Number')
        self.document.add_paragraph('            - nazwisko imię */ pełna nazwa klubu lub firmy: ').add_run(
            self.commissioning_body['name']).bold = True
        self.document.add_paragraph('            - adres: ').add_run(self.commissioning_body['address']).bold = True
        self.document.add_paragraph('            - tel: ').add_run(self.commissioning_body['tel']).bold = True
        self.document.add_paragraph('            - E-mail: ').add_run(self.commissioning_body['e-mail']).bold = True
        self.document.add_paragraph('            - NIP klubu/stowarzyszenia: ').add_run(
            self.commissioning_body['nip']).bold = True
        self.document.add_paragraph('Deklarowany okres i czas postoju: ', style='List Number')
        self.document.add_paragraph('            - od dnia: ').add_run(self.parking_peroid['from']).bold = True
        self.document.add_paragraph('            - do dnia: ').add_run(self.parking_peroid['to']).bold = True
        self.document.add_paragraph('')
        self.document.add_paragraph('Opłata za postój wynosi: ', style='List Number').add_run(
            str(self.fee['parking_fee'])).bold = True
        self.document.add_paragraph(
            'Czynsz płatny, zgodnie z wystawioną fakturą z góry, jednorazowo lub w czterech poniższych '
            'ratach za każdy kwartał do:')
        self.document.add_paragraph('      I. 15.05.2020, w kwocie: ').add_run(str(self.fee['quarter_fee'])).bold = True
        self.document.add_paragraph('     II. 15.08.2020, w kwocie: ').add_run(str(self.fee['quarter_fee'])).bold = True
        self.document.add_paragraph('    III. 15.11.2020, w kwocie: ').add_run(str(self.fee['quarter_fee'])).bold = True
        self.document.add_paragraph('    IV. 15.21.2020, w kwocie: ').add_run(str(self.fee['quarter_fee'])).bold = True
        self.document.add_paragraph('na rachunek Wynajmującego o nr  88 1030 1117 0000 0000 8899 5007.')
        self.document.add_paragraph('Adres do korespondencji: ', style='List Number').add_run(
            self.owner_details['address']).bold = True
        self.document.add_paragraph('Karta chipowa: ', style='List Number').add_run(self.chip_card).bold = True
        if self.media:
            p = self.document.add_paragraph('Dodatkowo zlecone usługi: ', style='List Number')
            p.add_run('184zł - miejsce postojowe na wodzie nieopomiarowane').bold = True
        self.document.add_paragraph('Miejsce postoju nr: ', style='List Number').add_run(self.parking_place).bold = True
        self.document.add_paragraph(
            '\nUprzejmie informujemy, że w trakcie postoju jachtu na przystani NCŻ AWFiS może zaistnieć '
            'konieczność zmiany miejsca lokalizacji postoju jachtu na inne zgodnie z ze wskazaniem '
            'upoważnionego pracownika przystani.\n\n'
            'Oświadczam, iż zapoznałem się z Regulaminem przystani Narodowego Centrum Żeglarstwa AWFiS '
            'Gdańsku, w tym zawarcia umowy na postój. W pełni go akceptuję co poświadczam własnoręcznym '
            'podpisem pod Deklaracją postoju. Oświadczam, iż w przypadku okresu postoju krótszego niż '
            'zadeklarowany (wynikający m.in. ze sprzedaży jachtu) mam świadomość, że stawka będzie '
            'rekalkulowana do krótszego okresu postoju, zgodnie z obowiązującym cennikiem.\n\n').paragraph_format.alignment = 3
        self.document.add_paragraph(".........................................................				 .........."
                                    "...............................................").add_run(
            '   podpis pracownika NCŻ                                                                                 '
            '             podpis (imię, nazwisko)').italic = True

        self.document.save('Deklaracja_{}.docx'.format(self.yacht['name']))



# yacht = {'name': 'Fordzik', 'registration_number': 'GD-102', 'home_port': 'Gdańsk', 'length': 12.00, 'width': 3.00,
#          'ytype': 'Motorowy'}
# fee = {'parking_fee': 8000, 'quarter_fee': 2000}
# owner_details = {'name': 'Adrian Bieliński', 'address': 'Anny Jagiellonki 23/14, Gdańsk'}
# parking_peroid = {'from': '01.05.2020', 'to': '31.10.2020'}
# commissioning_body = {'name': 'jw', 'address': 'jw',
#                       'tel': '510494063', 'e-mail': 'adimr52@gmail.com',
#                       'nip': 'BRAK'}
#
# document = CreateDeclaration('D-12', '15.02.2020', yacht, owner_details, commissioning_body, parking_peroid, fee, '1',
#                              True)
# document.create_document()

